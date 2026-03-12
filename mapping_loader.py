"""
mapping_loader.py
-----------------
Consomme le mapping.json produit par l'Annotateur de Rapport
et résout les valeurs depuis les fichiers Excel.

Usage:
    python mapping_loader.py mapping.json --excel-dir ./data
"""

import json, re, sys, argparse
from pathlib import Path

try:
    import openpyxl
except ImportError:
    print("pip install openpyxl")
    sys.exit(1)


def load_mapping(path: str) -> dict:
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def read_cell(wb_cache: dict, excel_dir: Path, filename: str, sheet: str, cell: str):
    key = filename
    if key not in wb_cache:
        wb_cache[key] = openpyxl.load_workbook(excel_dir / filename, data_only=True)
    return wb_cache[key][sheet][cell].value


def eval_formula(formula: str, wb_cache: dict, excel_dir: Path) -> float:
    """
    Résout une formule du type :
      (Budget_2024.xlsx|Résumé|B2 - Budget_2024.xlsx|Charges|C10) / Budget_2024.xlsx|Résumé|B2 * 100
    """
    ref_re = re.compile(r"([^|\s()+\-*/]+\.xlsx)\|([^|]+)\|([A-Z]+\d+)")

    def replacer(m):
        val = read_cell(wb_cache, excel_dir, m.group(1), m.group(2), m.group(3))
        return str(float(val))

    expression = ref_re.sub(replacer, formula)
    return eval(expression, {"__builtins__": {}})  # noqa: S307


def resolve_all(mapping_path: str, excel_dir: str = ".") -> dict[str, any]:
    data = load_mapping(mapping_path)
    excel_dir = Path(excel_dir)
    wb_cache = {}
    results = {}
    errors = {}

    for placeholder, cfg in data["placeholders"].items():
        if not cfg.get("_mapped"):
            results[placeholder] = None
            continue
        try:
            t = cfg["type"]
            if t == "cell":
                value = read_cell(wb_cache, excel_dir, cfg["source"], cfg["sheet"], cfg["cell"])
                # Fallback: si valeur vide/None et fallback défini
                if (value is None or str(value).strip() == "") and cfg.get("fallback"):
                    fb = cfg["fallback"]
                    value = read_cell(wb_cache, excel_dir, fb["source"], fb["sheet"], fb["cell"])
                results[placeholder] = value
            elif t == "formula":
                results[placeholder] = eval_formula(cfg["formula"], wb_cache, excel_dir)
            else:
                raise ValueError(f"Type inconnu: {t}")
        except Exception as e:
            errors[placeholder] = str(e)
            results[placeholder] = f"[ERREUR: {e}]"

    if errors:
        print(f"⚠  {len(errors)} erreur(s):", file=sys.stderr)
        for k, v in errors.items():
            print(f"   {k} → {v}", file=sys.stderr)

    return results


def fill_word(mapping_path: str, template: str, output: str, excel_dir: str = "."):
    """Remplit un template Word avec les valeurs résolues."""
    try:
        from docx import Document
    except ImportError:
        print("pip install python-docx")
        return

    values = resolve_all(mapping_path, excel_dir)
    doc = Document(template)

    def replace_para(para):
        for ph, val in values.items():
            if ph in para.text and val is not None:
                for run in para.runs:
                    if ph in run.text:
                        run.text = run.text.replace(ph, str(val))

    for para in doc.paragraphs:
        replace_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_para(para)

    doc.save(output)
    print(f"✓ Rapport généré : {output}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Résout les placeholders depuis mapping.json")
    parser.add_argument("mapping", help="Chemin vers mapping.json")
    parser.add_argument("--excel-dir", default=".", help="Dossier contenant les fichiers Excel")
    parser.add_argument("--template", help="Template Word (.docx)")
    parser.add_argument("--output", default="rapport_final.docx", help="Fichier Word de sortie")
    args = parser.parse_args()

    if args.template:
        fill_word(args.mapping, args.template, args.output, args.excel_dir)
    else:
        results = resolve_all(args.mapping, args.excel_dir)
        for ph, val in results.items():
            print(f"{ph:40s} → {val}")
